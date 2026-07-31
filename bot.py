# ===================================================================
#   TARJIMON BOT – MUKAMMAL (WEBHOOK + ADSGRAM + TARIF)
#            Moderatsiyadan o‘tish uchun optimallashtirilgan
# ===================================================================

import os
import logging
import sqlite3
import datetime
import threading
import time
from io import BytesIO
from flask import Flask, request, jsonify
import telebot
from telebot import types
from deep_translator import GoogleTranslator

# ---------------------- KONFIGURATSIYA ----------------------
BOT_TOKEN = os.environ.get("BOT_TOKEN")
ADMIN_ID = int(os.environ.get("ADMIN_ID", "123456789"))
WEBHOOK_URL = os.environ.get("WEBHOOK_URL")  # https://your-app.onrender.com/webhook
ADSGRAM_AD_UNIT_ID = os.environ.get("ADSGRAM_AD_UNIT_ID", "YOUR_AD_UNIT_ID")

if not BOT_TOKEN:
    raise ValueError("BOT_TOKEN environment variable is required!")
if not WEBHOOK_URL:
    raise ValueError("WEBHOOK_URL environment variable is required!")

# ---------------------- LOGGER ----------------------
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ---------------------- BOT, FLASK, TRANSLATOR ----------------------
bot = telebot.TeleBot(BOT_TOKEN, threaded=False)
translator = GoogleTranslator()
app = Flask(__name__)

# ---------------------- SQLITE ----------------------
DB_NAME = 'translator_bot.db'

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        user_id INTEGER PRIMARY KEY,
        lang TEXT DEFAULT 'uz',
        limit_count INTEGER DEFAULT 10,
        used_count INTEGER DEFAULT 0,
        total_ads_watched INTEGER DEFAULT 0,
        registered_at TEXT DEFAULT CURRENT_TIMESTAMP
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS stats (
        id INTEGER PRIMARY KEY CHECK (id = 1),
        total_users INTEGER DEFAULT 0,
        total_translations INTEGER DEFAULT 0,
        total_ads_watched INTEGER DEFAULT 0,
        daily_ads_count INTEGER DEFAULT 0,
        last_ad_reset TEXT DEFAULT CURRENT_DATE
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS settings (
        key TEXT PRIMARY KEY,
        value TEXT
    )''')
    
    c.execute("INSERT OR IGNORE INTO settings (key, value) VALUES ('daily_ad_limit', '10')")
    c.execute("INSERT OR IGNORE INTO settings (key, value) VALUES ('ad_reward_limit', '5')")
    c.execute("INSERT OR IGNORE INTO settings (key, value) VALUES ('free_limit', '7')")
    c.execute("INSERT OR IGNORE INTO settings (key, value) VALUES ('after_free', '4')")
    c.execute("INSERT OR IGNORE INTO settings (key, value) VALUES ('ad_count', '2')")
    
    c.execute("INSERT OR IGNORE INTO stats (id, total_users, total_translations, total_ads_watched, daily_ads_count, last_ad_reset) VALUES (1, 0, 0, 0, 0, date('now'))")
    
    c.execute('''CREATE TABLE IF NOT EXISTS daily_ads (
        user_id INTEGER,
        date TEXT,
        watched_count INTEGER DEFAULT 0,
        PRIMARY KEY (user_id, date)
    )''')
    
    conn.commit()
    conn.close()
    logger.info("Ma'lumotlar bazasi tayyorlandi.")

# ---------------------- DB YORDAMCHILARI ----------------------
def get_user(user_id):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT user_id, lang, limit_count, used_count, total_ads_watched FROM users WHERE user_id = ?", (user_id,))
    row = c.fetchone()
    conn.close()
    if row:
        return {
            'user_id': row[0],
            'lang': row[1],
            'limit': row[2],
            'used': row[3],
            'total_ads_watched': row[4]
        }
    return None

def create_or_update_user(user_id, lang='uz'):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("INSERT OR IGNORE INTO users (user_id, lang) VALUES (?, ?)", (user_id, lang))
    c.execute("UPDATE users SET lang = ? WHERE user_id = ?", (lang, user_id))
    conn.commit()
    conn.close()
    update_stats_total_users()

def get_user_lang_db(user_id):
    user = get_user(user_id)
    return user['lang'] if user else 'uz'

def get_setting(key):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT value FROM settings WHERE key = ?", (key,))
    row = c.fetchone()
    conn.close()
    return row[0] if row else None

def set_setting(key, value):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", (key, value))
    conn.commit()
    conn.close()

def get_tariff_settings():
    free_limit = int(get_setting('free_limit') or 7)
    after_free = int(get_setting('after_free') or 4)
    ad_count = int(get_setting('ad_count') or 2)
    return free_limit, after_free, ad_count

def get_daily_used(user_id):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT used_count FROM users WHERE user_id = ?", (user_id,))
    row = c.fetchone()
    conn.close()
    return row[0] if row else 0

def increment_daily_used(user_id):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("UPDATE users SET used_count = used_count + 1 WHERE user_id = ?", (user_id,))
    conn.commit()
    conn.close()

def reset_daily_used_all():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("UPDATE users SET used_count = 0")
    conn.commit()
    conn.close()
    logger.info("Kunlik used_count tiklandi.")

def get_daily_ads_watched(user_id):
    today = datetime.datetime.now().date().isoformat()
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT watched_count FROM daily_ads WHERE user_id = ? AND date = ?", (user_id, today))
    row = c.fetchone()
    conn.close()
    return row[0] if row else 0

def increment_daily_ads(user_id):
    today = datetime.datetime.now().date().isoformat()
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''INSERT INTO daily_ads (user_id, date, watched_count) 
                 VALUES (?, ?, 1) 
                 ON CONFLICT(user_id, date) DO UPDATE SET watched_count = watched_count + 1''', 
                 (user_id, today))
    conn.commit()
    conn.close()

def reset_daily_ads_all():
    today = datetime.datetime.now().date().isoformat()
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("DELETE FROM daily_ads WHERE date != ?", (today,))
    conn.commit()
    conn.close()

def can_translate(user_id):
    free_limit, after_free, ad_count = get_tariff_settings()
    used = get_daily_used(user_id)
    ads_watched = get_daily_ads_watched(user_id)

    if used < free_limit:
        return True, None

    pos = used - free_limit
    interval_index = pos // after_free
    required_ads = (interval_index + 1) * ad_count

    if ads_watched >= required_ads:
        return True, None
    else:
        need = required_ads - ads_watched
        return False, need

# ---------------------- STATISTIKA ----------------------
def get_stats():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT total_users, total_translations, total_ads_watched, daily_ads_count, last_ad_reset FROM stats WHERE id = 1")
    row = c.fetchone()
    conn.close()
    if row:
        return {
            'total_users': row[0],
            'total_translations': row[1],
            'total_ads_watched': row[2],
            'daily_ads_count': row[3],
            'last_ad_reset': row[4]
        }
    return None

def update_stats_total_users():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM users")
    count = c.fetchone()[0]
    c.execute("UPDATE stats SET total_users = ? WHERE id = 1", (count,))
    conn.commit()
    conn.close()

def increment_daily_ads_stat():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("UPDATE stats SET daily_ads_count = daily_ads_count + 1 WHERE id = 1")
    conn.commit()
    conn.close()

def reset_daily_ads_stat():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("UPDATE stats SET daily_ads_count = 0, last_ad_reset = date('now') WHERE id = 1")
    conn.commit()
    conn.close()

# ---------------------- AVTOMATIK RESET ----------------------
def daily_reset():
    reset_daily_used_all()
    reset_daily_ads_all()
    reset_daily_ads_stat()
    logger.info("Kunlik reset bajarildi.")

def check_daily_reset():
    stats = get_stats()
    if stats:
        today = datetime.datetime.now().date().isoformat()
        if stats['last_ad_reset'] != today:
            daily_reset()
    threading.Timer(3600, check_daily_reset).start()

# ---------------------- TILGA MOS MATNLAR ----------------------
def get_account_text(user_id, lang):
    user = get_user(user_id)
    if not user:
        return "❌ Foydalanuvchi topilmadi."
    used = get_daily_used(user_id)
    ads = get_daily_ads_watched(user_id)
    total_ads = user.get('total_ads_watched', 0)
    free_limit, after_free, ad_count = get_tariff_settings()
    
    texts = {
        'uz': f"👤 Hisobingiz:\nBugun tarjima: {used}\nBepul limit: {free_limit}\nOraliq: {after_free}\nReklama soni: {ad_count}\nBugun reklama: {ads}\nJami reklama: {total_ads}",
        'ru': f"👤 Ваш аккаунт:\nСегодня переводов: {used}\nБесплатный лимит: {free_limit}\nИнтервал: {after_free}\nКол-во реклам: {ad_count}\nСегодня реклам: {ads}\nВсего реклам: {total_ads}",
        'en': f"👤 Your account:\nToday translations: {used}\nFree limit: {free_limit}\nInterval: {after_free}\nAd count: {ad_count}\nToday ads: {ads}\nTotal ads: {total_ads}"
    }
    return texts.get(lang, texts['uz'])

def get_limit_text(user_id, lang):
    user = get_user(user_id)
    if not user:
        return "❌ Foydalanuvchi topilmadi."
    used = get_daily_used(user_id)
    free_limit, after_free, ad_count = get_tariff_settings()
    ads_watched = get_daily_ads_watched(user_id)
    
    texts = {
        'uz': f"💰 Bugungi tarjimalar: {used}\n📊 Bepul limit: {free_limit}\n📈 Oraliq: har {after_free} ta tarjimadan keyin {ad_count} ta reklama\n🎥 Bugungi reklamalar: {ads_watched}",
        'ru': f"💰 Сегодня переводов: {used}\n📊 Бесплатный лимит: {free_limit}\n📈 Интервал: после каждых {after_free} переводов {ad_count} рекламы\n🎥 Сегодня реклам: {ads_watched}",
        'en': f"💰 Today translations: {used}\n📊 Free limit: {free_limit}\n📈 Interval: after every {after_free} translations {ad_count} ads\n🎥 Today ads: {ads_watched}"
    }
    return texts.get(lang, texts['uz'])

# ---------------------- FAYL YARATISH ----------------------
def create_translation_file(text, format_type='txt'):
    if format_type == 'txt':
        file_io = BytesIO()
        file_io.write(text.encode('utf-8'))
        file_io.seek(0)
        return file_io, 'txt'
    elif format_type == 'docx':
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph(text)
            file_io = BytesIO()
            doc.save(file_io)
            file_io.seek(0)
            return file_io, 'docx'
        except ImportError:
            return create_translation_file(text, 'txt')
    elif format_type == 'pdf':
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            file_io = BytesIO()
            doc = SimpleDocTemplate(file_io, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            for line in text.split('\n'):
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 12))
            doc.build(story)
            file_io.seek(0)
            return file_io, 'pdf'
        except ImportError:
            return create_translation_file(text, 'txt')
    else:
        return create_translation_file(text, 'txt')

# ---------------------- KEYBOARDLAR ----------------------
def get_main_menu_keyboard(lang='uz', user_id=None):
    texts = {
        'uz': {'translate': "📝 Tarjima", 'tariffs': "📊 Tariflar", 'limit': "💰 Limit", 'ad': "📢 Reklama", 'language': "🌐 Til", 'account': "👤 Hisob", 'admin': "🔐 Admin"},
        'ru': {'translate': "📝 Перевод", 'tariffs': "📊 Тарифы", 'limit': "💰 Лимит", 'ad': "📢 Реклама", 'language': "🌐 Язык", 'account': "👤 Аккаунт", 'admin': "🔐 Админ"},
        'en': {'translate': "📝 Translate", 'tariffs': "📊 Tariffs", 'limit': "💰 Limit", 'ad': "📢 Ad", 'language': "🌐 Language", 'account': "👤 Account", 'admin': "🔐 Admin"}
    }
    t = texts.get(lang, texts['uz'])
    keyboard = types.InlineKeyboardMarkup(row_width=2)
    btns = [
        types.InlineKeyboardButton(t['translate'], callback_data="translate"),
        types.InlineKeyboardButton(t['tariffs'], callback_data="tariffs"),
        types.InlineKeyboardButton(t['limit'], callback_data="limit"),
        types.InlineKeyboardButton(t['ad'], callback_data="ad"),
        types.InlineKeyboardButton(t['language'], callback_data="language"),
        types.InlineKeyboardButton(t['account'], callback_data="account")
    ]
    if user_id == ADMIN_ID:
        btns.append(types.InlineKeyboardButton(t['admin'], callback_data="admin_panel"))
    keyboard.add(*btns)
    return keyboard

def get_language_keyboard():
    keyboard = types.InlineKeyboardMarkup(row_width=3)
    keyboard.add(
        types.InlineKeyboardButton("🇺🇿 O'zbek", callback_data="lang_uz"),
        types.InlineKeyboardButton("🇷🇺 Русский", callback_data="lang_ru"),
        types.InlineKeyboardButton("🇬🇧 English", callback_data="lang_en")
    )
    return keyboard

def get_translation_languages_keyboard():
    keyboard = types.InlineKeyboardMarkup(row_width=3)
    buttons = [
        ("🇬🇧 Ingliz", "trans_en"), ("🇷🇺 Rus", "trans_ru"), ("🇺🇿 O'zbek", "trans_uz"),
        ("🇹🇷 Turk", "trans_tr"), ("🇫🇷 Fransuz", "trans_fr"), ("🇩🇪 Nemis", "trans_de"),
        ("🇪🇸 Ispan", "trans_es"), ("🇨🇳 Xitoy", "trans_zh-cn"), ("🇦🇪 Arab", "trans_ar")
    ]
    for label, callback in buttons:
        keyboard.add(types.InlineKeyboardButton(label, callback_data=callback))
    return keyboard

def get_result_format_keyboard(lang='uz'):
    texts = {'uz': {"text": "📝 Matn", "file": "📄 Fayl"}, 'ru': {"text": "📝 Текст", "file": "📄 Файл"}, 'en': {"text": "📝 Text", "file": "📄 File"}}
    t = texts.get(lang, texts['uz'])
    keyboard = types.InlineKeyboardMarkup(row_width=2)
    keyboard.add(
        types.InlineKeyboardButton(t['text'], callback_data="result_text"),
        types.InlineKeyboardButton(t['file'], callback_data="result_file")
    )
    return keyboard

def get_ad_watch_keyboard(lang='uz'):
    texts = {'uz': "🎥 Reklama ko'rish", 'ru': "🎥 Посмотреть рекламу", 'en': "🎥 Watch ad"}
    keyboard = types.InlineKeyboardMarkup()
    keyboard.add(types.InlineKeyboardButton(texts.get(lang, texts['uz']), callback_data="watch_ad"))
    return keyboard

def get_admin_keyboard(lang='uz'):
    texts = {
        'uz': {'stats': "📊 Statistika", 'ad_settings': "📢 Reklama sozlamalari", 'send_ad': "📨 Reklama yuborish", 'tariff_settings': "⚙️ Tarif sozlamalari", 'reset_daily': "🔄 Kunlik reklamani tiklash", 'back': "🔙 Orqaga"},
        'ru': {'stats': "📊 Статистика", 'ad_settings': "📢 Настройки рекламы", 'send_ad': "📨 Отправить рекламу", 'tariff_settings': "⚙️ Настройки тарифов", 'reset_daily': "🔄 Сбросить дневную рекламу", 'back': "🔙 Назад"},
        'en': {'stats': "📊 Statistics", 'ad_settings': "📢 Ad settings", 'send_ad': "📨 Send ad", 'tariff_settings': "⚙️ Tariff settings", 'reset_daily': "🔄 Reset daily ads", 'back': "🔙 Back"}
    }
    t = texts.get(lang, texts['uz'])
    keyboard = types.InlineKeyboardMarkup(row_width=2)
    keyboard.add(
        types.InlineKeyboardButton(t['stats'], callback_data="admin_stats"),
        types.InlineKeyboardButton(t['ad_settings'], callback_data="admin_ad_settings"),
        types.InlineKeyboardButton(t['send_ad'], callback_data="admin_send_ad"),
        types.InlineKeyboardButton(t['tariff_settings'], callback_data="admin_tariff_settings"),
        types.InlineKeyboardButton(t['reset_daily'], callback_data="admin_reset_daily"),
        types.InlineKeyboardButton(t['back'], callback_data="admin_back")
    )
    return keyboard

def get_tariff_settings_keyboard(lang='uz'):
    texts = {'uz': {'change': "✏️ O'zgartirish", 'reset': "🔄 Default", 'back': "🔙 Orqaga"},
             'ru': {'change': "✏️ Изменить", 'reset': "🔄 Сброс", 'back': "🔙 Назад"},
             'en': {'change': "✏️ Change", 'reset': "🔄 Reset", 'back': "🔙 Back"}}
    t = texts.get(lang, texts['uz'])
    keyboard = types.InlineKeyboardMarkup(row_width=2)
    keyboard.add(
        types.InlineKeyboardButton(t['change'], callback_data="tariff_change"),
        types.InlineKeyboardButton(t['reset'], callback_data="tariff_reset"),
        types.InlineKeyboardButton(t['back'], callback_data="tariff_back")
    )
    return keyboard

# ---------------------- TEMP DATA ----------------------
temp_data = {}

# ---------------------- BOT HANDLERLAR ----------------------
@bot.message_handler(commands=['start'])
def start_command(message):
    user_id = message.from_user.id
    user = get_user(user_id)
    if not user:
        create_or_update_user(user_id, 'uz')
        lang = 'uz'
    else:
        lang = user['lang']
    bot.send_message(
        message.chat.id,
        "Assalomu alaykum! Men tarjimon botman. Quyidagi menyudan foydalaning:",
        reply_markup=get_main_menu_keyboard(lang, user_id)
    )
    logger.info(f"User {user_id} started bot")

@bot.message_handler(commands=['admin'])
def admin_command(message):
    if message.from_user.id != ADMIN_ID:
        bot.reply_to(message, "⛔ Siz admin emassiz!")
        return
    lang = get_user_lang_db(message.from_user.id)
    bot.send_message(message.chat.id, "🔐 Admin paneli:", reply_markup=get_admin_keyboard(lang))

@bot.message_handler(func=lambda message: True, content_types=['text'])
def handle_text(message):
    user_id = message.from_user.id
    text = message.text
    if text.startswith('/'):
        return
    user = get_user(user_id)
    if not user:
        create_or_update_user(user_id, 'uz')
        lang = 'uz'
        bot.send_message(message.chat.id, "Assalomu alaykum! Men tarjimon botman. Quyidagi menyudan foydalaning:", reply_markup=get_main_menu_keyboard(lang, user_id))
        return
    lang = user['lang']
    if user_id not in temp_data:
        temp_data[user_id] = {}
    temp_data[user_id]['temp_text'] = text
    temp_data[user_id]['is_file'] = False
    prompt = {'uz': "📝 Quyidagi matn tarjima qilinadi:\n\n", 'ru': "📝 Следующий текст будет переведён:\n\n", 'en': "📝 The following text will be translated:\n\n"}.get(lang, "📝 Quyidagi matn tarjima qilinadi:\n\n")
    bot.reply_to(message, prompt + text[:500] + ("..." if len(text) > 500 else ""), reply_markup=get_translation_languages_keyboard())

@bot.message_handler(content_types=['document'])
def handle_document(message):
    user_id = message.from_user.id
    lang = get_user_lang_db(user_id)
    file_info = bot.get_file(message.document.file_id)
    file_name = message.document.file_name
    if not file_name.lower().endswith(('.txt', '.docx', '.pdf')):
        bot.reply_to(message, "❌ Faqat .txt, .docx va .pdf formatdagi fayllarni qabul qilaman.")
        return
    try:
        data = bot.download_file(file_info.file_path)
        if file_name.endswith('.txt'):
            content = data.decode('utf-8')
        elif file_name.endswith('.docx'):
            from docx import Document
            doc = Document(BytesIO(data))
            content = '\n'.join([p.text for p in doc.paragraphs])
        elif file_name.endswith('.pdf'):
            import PyPDF2
            reader = PyPDF2.PdfReader(BytesIO(data))
            content = ''.join([page.extract_text() or '' for page in reader.pages])
        else:
            bot.reply_to(message, "❌ Qo'llab-quvvatlanmaydigan format.")
            return
    except Exception as e:
        logger.error(f"Faylni o'qish xatosi: {e}")
        bot.reply_to(message, "❌ Faylni o'qib bo'lmadi.")
        return
    if user_id not in temp_data:
        temp_data[user_id] = {}
    temp_data[user_id]['temp_text'] = content
    temp_data[user_id]['is_file'] = True
    prompt = {'uz': f"📄 Fayl qabul qilindi: {file_name}\n({len(content)} belgi) tarjima qilinadi:", 'ru': f"📄 Файл получен: {file_name}\n({len(content)} символов) будет переведён:", 'en': f"📄 File received: {file_name}\n({len(content)} characters) will be translated:"}.get(lang, f"📄 Fayl qabul qilindi: {file_name}\n({len(content)} belgi) tarjima qilinadi:")
    bot.reply_to(message, prompt + "\n\n" + content[:300] + ("..." if len(content) > 300 else ""), reply_markup=get_translation_languages_keyboard())

# ---------------------- CALLBACK HANDLER ----------------------
@bot.callback_query_handler(func=lambda call: True)
def handle_callback(call):
    user_id = call.from_user.id
    lang = get_user_lang_db(user_id)
    data = call.data

    if data.startswith("lang_"):
        new_lang = data.split("_")[1]
        create_or_update_user(user_id, new_lang)
        bot.answer_callback_query(call.id, text=f"Til o'zgartirildi: {new_lang.upper()}")
        bot.edit_message_text("Til o'zgartirildi. Quyidagi menyudan foydalaning:", chat_id=call.message.chat.id, message_id=call.message.message_id, reply_markup=get_main_menu_keyboard(new_lang, user_id))
        return

    if data == "limit":
        bot.send_message(call.message.chat.id, get_limit_text(user_id, lang))
        bot.answer_callback_query(call.id)
        return

    if data == "account":
        bot.send_message(call.message.chat.id, get_account_text(user_id, lang))
        bot.answer_callback_query(call.id)
        return

    if data == "translate":
        bot.answer_callback_query(call.id)
        bot.send_message(call.message.chat.id, "📝 Tarjima qilish uchun matn yoki fayl yuboring.")
        return

    if data == "tariffs":
        bot.answer_callback_query(call.id)
        free_limit, after_free, ad_count = get_tariff_settings()
        texts = {'uz': f"📊 Joriy tarif:\n- Bepul limit: {free_limit}\n- Har {after_free} ta tarjimadan keyin {ad_count} ta reklama", 'ru': f"📊 Текущий тариф:\n- Бесплатно: {free_limit}\n- После {after_free} переводов {ad_count} рекламы", 'en': f"📊 Current tariff:\n- Free: {free_limit}\n- After {after_free} translations {ad_count} ads"}
        bot.send_message(call.message.chat.id, texts.get(lang, texts['uz']))
        return

    if data == "ad":
        bot.answer_callback_query(call.id)
        bot.send_message(call.message.chat.id, "📢 Reklama ko'rib limit oshiring:", reply_markup=get_ad_watch_keyboard(lang))
        return

    if data == "language":
        bot.answer_callback_query(call.id)
        bot.edit_message_text("🌐 Tilni tanlang:", chat_id=call.message.chat.id, message_id=call.message.message_id, reply_markup=get_language_keyboard())
        return

    if data == "watch_ad":
        stats = get_stats()
        if not stats:
            bot.answer_callback_query(call.id, text="❌ Xatolik.")
            return
        daily_limit = int(get_setting('daily_ad_limit') or 10)
        if stats['daily_ads_count'] >= daily_limit:
            bot.answer_callback_query(call.id, text="❌ Bugungi reklama limiti tugagan.")
            return
        ad_link = f"https://t.me/adsgram_bot?start=reward_{ADSGRAM_AD_UNIT_ID}_{user_id}"
        bot.send_message(call.message.chat.id, f"🎥 Reklamani ko'rish uchun bosing:\n👉 {ad_link}", reply_markup=types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("📺 Reklamani ko'rish", url=ad_link)))
        if user_id not in temp_data:
            temp_data[user_id] = {}
        temp_data[user_id]['waiting_ad'] = True
        bot.answer_callback_query(call.id, text="Havola yuborildi.")
        return

    if data.startswith('trans_'):
        target_lang = data.split('_')[1]
        user_temp = temp_data.get(user_id, {})
        text_to_translate = user_temp.get('temp_text')
        if not text_to_translate:
            bot.answer_callback_query(call.id, text="❌ Matn topilmadi.")
            return
        allowed, need_ads = can_translate(user_id)
        if not allowed:
            bot.answer_callback_query(call.id, text=f"❌ {need_ads} ta reklama kerak.")
            msgs = {'uz': f"⚠️ {need_ads} ta reklama ko'ring.", 'ru': f"⚠️ Посмотрите {need_ads} рекламы.", 'en': f"⚠️ Watch {need_ads} ads."}
            bot.send_message(call.message.chat.id, msgs.get(lang, msgs['uz']), reply_markup=get_ad_watch_keyboard(lang))
            return
        try:
            translated = translator.translate(text_to_translate, target=target_lang)
        except Exception as e:
            logger.error(f"Tarjima xatosi: {e}")
            bot.answer_callback_query(call.id, text="❌ Xatolik.")
            return
        increment_daily_used(user_id)
        if user_id not in temp_data:
            temp_data[user_id] = {}
        temp_data[user_id]['translated_text'] = translated
        temp_data[user_id]['target_lang'] = target_lang
        bot.answer_callback_query(call.id, text="✅ Tayyor! Formatni tanlang.")
        bot.send_message(call.message.chat.id, "📥 Formatni tanlang:", reply_markup=get_result_format_keyboard(lang))
        return

    if data.startswith('result_'):
        result_type = data.split('_')[1]
        user_temp = temp_data.get(user_id, {})
        translated = user_temp.get('translated_text')
        target_lang = user_temp.get('target_lang', 'en')
        if not translated:
            bot.answer_callback_query(call.id, text="❌ Matn topilmadi.")
            return
        if result_type == 'text':
            bot.send_message(call.message.chat.id, f"🌐 Tarjima ({target_lang.upper()}):\n\n{translated}", reply_markup=get_main_menu_keyboard(lang, user_id))
            bot.answer_callback_query(call.id, text="✅ Matn jo'natildi.")
        else:
            file_io, ext = create_translation_file(translated, 'txt')
            try:
                bot.send_document(call.message.chat.id, file_io, caption=f"📄 Tarjima fayli ({target_lang.upper()}).{ext}", reply_markup=get_main_menu_keyboard(lang, user_id))
                bot.answer_callback_query(call.id, text="✅ Fayl jo'natildi.")
            except Exception as e:
                logger.error(f"Fayl xatosi: {e}")
                bot.answer_callback_query(call.id, text="❌ Xatolik.")
            finally:
                file_io.close()
        temp_data[user_id].pop('translated_text', None)
        temp_data[user_id].pop('target_lang', None)
        temp_data[user_id].pop('temp_text', None)
        return

    # Admin paneli
    if user_id != ADMIN_ID:
        bot.answer_callback_query(call.id, text="⛔ Ruxsat yo'q!")
        return

    if data == "admin_panel":
        bot.answer_callback_query(call.id)
        bot.send_message(call.message.chat.id, "🔐 Admin paneli:", reply_markup=get_admin_keyboard(lang))
        return
    elif data == "admin_stats":
        bot.answer_callback_query(call.id)
        show_statistics(call.message, lang)
        return
    elif data == "admin_ad_settings":
        bot.answer_callback_query(call.id)
        show_ad_settings(call.message, lang)
        return
    elif data == "admin_send_ad":
        bot.answer_callback_query(call.id)
        msg = bot.send_message(call.message.chat.id, "📨 Reklama matni:")
        bot.register_next_step_handler(msg, admin_send_ad_step)
        return
    elif data == "admin_tariff_settings":
        bot.answer_callback_query(call.id)
        free_limit, after_free, ad_count = get_tariff_settings()
        text = f"📊 Joriy: Bepul={free_limit}, Oraliq={after_free}, Reklama={ad_count}"
        bot.send_message(call.message.chat.id, text, reply_markup=get_tariff_settings_keyboard(lang))
        return
    elif data == "admin_reset_daily":
        daily_reset()
        bot.answer_callback_query(call.id, text="✅ Tiklandi!")
        return
    elif data == "admin_back":
        bot.answer_callback_query(call.id)
        bot.edit_message_text("🔐 Admin paneli:", chat_id=call.message.chat.id, message_id=call.message.message_id, reply_markup=get_admin_keyboard(lang))
        return

    # Tarif sozlamalari
    if data == "tariff_change":
        bot.answer_callback_query(call.id)
        msg = bot.send_message(call.message.chat.id, "✏️ 3 ta son yozing (masalan: `7 4 2`):")
        bot.register_next_step_handler(msg, tariff_change_step)
        return
    elif data == "tariff_reset":
        set_setting('free_limit', '7'); set_setting('after_free', '4'); set_setting('ad_count', '2')
        bot.answer_callback_query(call.id, text="✅ Default qaytarildi!")
        bot.edit_message_text("📊 Default qaytarildi (7,4,2).", chat_id=call.message.chat.id, message_id=call.message.message_id, reply_markup=get_tariff_settings_keyboard(lang))
        return
    elif data == "tariff_back":
        bot.answer_callback_query(call.id)
        bot.edit_message_text("🔐 Admin paneli:", chat_id=call.message.chat.id, message_id=call.message.message_id, reply_markup=get_admin_keyboard(lang))
        return

# ---------------------- STEP FUNKSIYALAR ----------------------
def tariff_change_step(message):
    if message.from_user.id != ADMIN_ID:
        return
    try:
        parts = message.text.split()
        if len(parts) != 3:
            raise ValueError
        fl, af, ac = int(parts[0]), int(parts[1]), int(parts[2])
        if fl < 0 or af < 1 or ac < 1:
            raise ValueError
        set_setting('free_limit', str(fl)); set_setting('after_free', str(af)); set_setting('ad_count', str(ac))
        bot.reply_to(message, f"✅ Yangilandi: Bepul={fl}, Oraliq={af}, Reklama={ac}")
    except:
        bot.reply_to(message, "❌ Xato! Masalan: `7 4 2`")

def admin_send_ad_step(message):
    if message.from_user.id != ADMIN_ID:
        return
    ad_text = message.text
    sent = 0
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT user_id FROM users")
    users = c.fetchall()
    conn.close()
    for (uid,) in users:
        try:
            bot.send_message(uid, f"📢 Reklama:\n\n{ad_text}")
            sent += 1
            time.sleep(0.05)
        except:
            pass
    bot.reply_to(message, f"✅ {sent} ta foydalanuvchiga yuborildi.")

def show_statistics(message, lang):
    stats = get_stats()
    if not stats:
        bot.send_message(message.chat.id, "❌ Statistika mavjud emas.")
        return
    daily_limit = get_setting('daily_ad_limit') or 10
    free_limit, after_free, ad_count = get_tariff_settings()
    texts = {
        'uz': f"📊 **Statistika**:\n👥 {stats['total_users']}\n📝 {stats['total_translations']}\n🎥 {stats['total_ads_watched']}\n📅 Kunlik reklama limiti: {daily_limit}\n📢 Bugun: {stats['daily_ads_count']}\n⚙️ Tarif: Bepul={free_limit}, Oraliq={after_free}, Reklama={ad_count}",
        'ru': f"📊 **Статистика**:\n👥 {stats['total_users']}\n📝 {stats['total_translations']}\n🎥 {stats['total_ads_watched']}\n📅 Дневной лимит: {daily_limit}\n📢 Сегодня: {stats['daily_ads_count']}\n⚙️ Тариф: Бесплатно={free_limit}, Интервал={after_free}, Реклам={ad_count}",
        'en': f"📊 **Statistics**:\n👥 {stats['total_users']}\n📝 {stats['total_translations']}\n🎥 {stats['total_ads_watched']}\n📅 Daily ad limit: {daily_limit}\n📢 Today: {stats['daily_ads_count']}\n⚙️ Tariff: Free={free_limit}, Interval={after_free}, Ads={ad_count}"
    }
    bot.send_message(message.chat.id, texts.get(lang, texts['uz']), parse_mode='Markdown')

def show_ad_settings(message, lang):
    daily_limit = get_setting('daily_ad_limit') or 10
    stats = get_stats()
    daily_ads = stats['daily_ads_count'] if stats else 0
    texts = {
        'uz': f"📢 Kunlik reklama limiti: {daily_limit}\n📊 Bugun: {daily_ads}",
        'ru': f"📢 Дневной лимит рекламы: {daily_limit}\n📊 Сегодня: {daily_ads}",
        'en': f"📢 Daily ad limit: {daily_limit}\n📊 Today: {daily_ads}"
    }
    bot.send_message(message.chat.id, texts.get(lang, texts['uz']))

# ---------------------- ADSGRAM CALLBACK (MUHIM) ----------------------
@app.route('/adsgram_callback', methods=['GET', 'POST'])
def adsgram_callback():
    # GET so'rov – moderator tekshiruvi uchun
    if request.method == 'GET':
        logger.info("Adsgram callback GET request received (health check)")
        return "Adsgram callback endpoint is active and working!", 200
    
    # POST so'rov – haqiqiy mukofot
    try:
        data = request.get_json()
        logger.info(f"Adsgram callback DATA: {data}")
        if not data:
            return "Invalid JSON", 400
        
        user_id = data.get('user_id')
        status = data.get('status')
        reward = data.get('reward', 5)
        
        if not user_id:
            return "Missing user_id", 400
        
        logger.info(f"Adsgram: user={user_id}, status={status}, reward={reward}")
        
        if status == 'completed':
            user = get_user(user_id)
            if user:
                increment_daily_ads(user_id)
                increment_daily_ads_stat()
                try:
                    lang = get_user_lang_db(user_id)
                    bot.send_message(
                        user_id,
                        "🎉 Reklama muvaffaqiyatli ko'rildi! Endi tarjima qilishingiz mumkin.",
                        reply_markup=get_main_menu_keyboard(lang, user_id)
                    )
                except Exception as e:
                    logger.error(f"Xabar yuborish xatosi: {e}")
            else:
                logger.warning(f"User {user_id} topilmadi.")
        else:
            try:
                bot.send_message(user_id, "❌ Reklama bekor qilindi yoki xatolik.")
            except:
                pass
        
        return "OK", 200
    except Exception as e:
        logger.error(f"Adsgram callback xatosi: {e}")
        return "Error", 500

# ---------------------- WEBHOOK ----------------------
@app.route('/webhook', methods=['POST'])
def webhook():
    if request.headers.get('content-type') == 'application/json':
        json_string = request.get_data().decode('utf-8')
        update = telebot.types.Update.de_json(json_string)
        bot.process_new_updates([update])
        return '', 200
    return 'Unsupported content type', 400

@app.route('/')
def index():
    return "Translator bot is running! Webhook is active."

def set_webhook():
    bot.remove_webhook()
    bot.set_webhook(url=WEBHOOK_URL)
    logger.info(f"Webhook set to {WEBHOOK_URL}")

# ---------------------- ISHGA TUSHIRISH ----------------------
if __name__ == '__main__':
    init_db()
    check_daily_reset()
    set_webhook()
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
